import docx
from docx.shared import Inches
from pathlib import Path

# Change this path to where you want the CV saved
path = Path(Path.home().joinpath("Desktop", "CSV", "jeremy_updated.docx"))
path.parent.mkdir(parents=True, exist_ok=True)

# Personal details (keep these accurate)
FULL_NAME = "MULWA KITALA JEREMIAH"
ADDRESS_LINES = [
    "Meru University of Science and Technology",
    "P.O Box 972-60200",
    "MERU"
]
EMAIL = "mulwajeremy24@gmail.com"
PHONE = "0713916894"
GITHUB = "https://github.com/jeremy-dhy98"  # <- replace with your GitHub
LINKEDIN = "https://www.linkedin.com/in/jeremiah-kitala-15aa57288"  # <- replace with your LinkedIn

# Photo (optional) — change the filename or set to None to skip
img_path = "IMG_20250208_202537.jpg"

# Build document
doc = docx.Document()

# Header / contact
header = doc.add_paragraph()
header_run = header.add_run(FULL_NAME + "\n")
header_run.bold = True
for line in ADDRESS_LINES:
    header.add_run(line + "\n")
header.add_run(f"Email: {EMAIL}\n")
header.add_run(f"Phone: {PHONE}\n")
header.add_run(f"GitHub: {GITHUB}\n")
header.add_run(f"LinkedIn: {LINKEDIN}\n")

# Professional summary (stronger than generic objective)
doc.add_heading("Professional Summary", level=1)
doc.add_paragraph(
    (
        "Applied Statistics graduate with strong foundations in mathematics and hands-on experience in data analysis and data engineering. "
        "Proficient in Python for data cleaning, analysis and automation (pandas, numpy, scikit-learn), experienced with SQL and database management, and competent in C# for software tasks. "
        "Skilled at building end-to-end data pipelines, performing exploratory data analysis and statistical modelling, and producing clear visualisations and reports to support decisions. "
        "Eager to contribute technical skill and analytical thinking to a data-driven team."
    )
)

# Education
doc.add_heading("Education", level=1)
doc.add_paragraph("(2021 - 2025)\nBachelor of Science in Mathematics and Computer Science — Specialization: Applied Statistics\nMeru University of Science and Technology")

# Relevant coursework (short list)
doc.add_heading("Relevant Coursework", level=2)
for course in [
    "Probability & Mathematical Statistics",
    "Statistical Inference and Regression",
    "Time Series Analysis",
    "Calculus and Linear Algebra",
    "Database Systems and SQL",
    "Algorithms and Data Structures"
]:
    doc.add_paragraph(course, style='List Bullet')

# Technical skills
doc.add_heading("Technical Skills", level=1)
# Categorized skills make scan-ability better
skills = {
    "Programming": "Python (pandas, numpy), C#",
    "Data Analysis / ML": "pandas, numpy, scikit-learn, statsmodels, matplotlib, seaborn",
    "Data Engineering / Databases": "SQL (querying & schema design), SQLite, PostgreSQL (familiar), ETL pipelines, API integration (REST, WebSockets)",
    "Statistics & Math": "Hypothesis testing, regression, time-series, probability, calculus, linear algebra",
    "Tools / Other": "Git, Jupyter, MATLAB, SPSS, STATA, Microsoft Excel"
}
for k, v in skills.items():
    p = doc.add_paragraph()
    p.add_run(f"{k}: ").bold = True
    p.add_run(v)

# Experience / Projects
doc.add_heading("Projects & Experience", level=1)

# Project 1: Data analysis
p1 = doc.add_paragraph()
p1.add_run("Data Analysis Project — Exploratory Data Analysis (EDA)").bold = True
p1.add_run("\n• Conducted EDA to identify trends, outliers and key metrics; cleaned and transformed datasets using pandas; produced visualisations and summary reports used for decision-making.\n")

# Project 2: Data engineering / trading pipeline (shows applied engineering skills)
p2 = doc.add_paragraph()
p2.add_run("Automated Data Pipeline & Backtesting (Personal/Academic)").bold = True
p2.add_run(
    "\n• Built end-to-end pipelines to fetch historical and real-time cryptocurrency data (APIs & WebSockets). Implemented indicators (RSI, MACD, ADX) and walked data through backtesting loops; improved WebSocket stability with reconnection logic and logging for reproducible results."
)

# Project 3: Software / Tools
p3 = doc.add_paragraph()
p3.add_run("Tooling & Scripts for Analysis").bold = True
p3.add_run(
    "\n• Developed scripts and utilities in Python and C# for data extraction, transformation and reporting. Automated repetitive tasks to save time and reduce manual error."
)

# Achievements & Certifications
doc.add_heading("Achievements & Certifications", level=1)
doc.add_paragraph("• Certificate — SPSS & STATA course completed at Kesap Research Center")
doc.add_paragraph("• Active contributor to university innovation initiatives (Meru Innovation Club)")

# Extracurricular
doc.add_heading("Extracurricular Activities", level=1)
doc.add_paragraph("• Member — Meru Innovation Club: participated in workshops and collaborative projects related to applied computing and entrepreneurship.")

# References
doc.add_heading("References", level=1)
doc.add_paragraph("Mrs. Christine Gacheri\nH.O.D Mathematics\nEmail: cmutuura@must.ac.ke\nPhone: 0723674987")

# Add photo if available
try:
    if img_path:
        doc.add_picture(img_path, width=Inches(1.5))
except Exception as e:
    # If photo not found, we continue without stopping
    doc.add_paragraph("(Photo omitted — ensure img_path points to a valid image if you want a photo included.)")

# Save
doc.save(path)
print(f"Saved updated CV to: {path}")
