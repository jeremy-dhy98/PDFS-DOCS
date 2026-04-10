import docx
from docx.shared import Inches
from pathlib import Path
from invitatin_cards import invites

doc = docx.Document()
doc.add_heading("Graduation invites", 0) # Add heading

para1 = str(invites[0])
doc.add_paragraph(para1).bold=True # add a paragraph
doc.paragraphs[0].runs[0].add_break() # line break
# doc.add_page_break()

para2 = str(invites[1])
doc.add_paragraph(para2). italic=True # another paragraph

para3 = str(invites[2])
doc.add_page_break()
doc.add_paragraph(para3)
doc.paragraphs[0].runs[0].add_break()

para4 = str(invites[3])
doc.add_paragraph(para4)

img_path_ =  "er.jpg"
doc.add_picture(img_path_, width=Inches(3), height=Inches(4))

path = Path(Path.home().joinpath("Desktop","CSV","graduation_invites.docs"))
doc.save(path)
