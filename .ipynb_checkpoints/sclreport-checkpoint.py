import docx
from docx.shared import Inches
from pathlib import Path

path = Path(Path.home().joinpath("Desktop", "CSV", "AttachReport.docx"))

doc = docx.Document()

# Cover page (Heading level 0 for the title)
doc.add_heading(text="""MERU UNIVERSITY OF SCIENCE AND TECHNOLOGY 
SCHOOL OF PURE AND APPLIED SCIENCES (SPAS)
ATTACHMENT REPORT""", level=0)
doc.add_paragraph(text="""NAME: MULWA KITALA JEREMIAH
REG.NO: SC205/103633/20
COURSE: Bachelor of Science in Mathematics and Computer Science (STATISTICS OPTION)
YEAR: III
DPT: MATHEMATICS
ORGANISATION: SAFRA DATA
C.E.O: NG'ANG'A NGARUIYA
SUPERVISOR: MR. BENARD WENJE
DURATION: III Months
COMMENCED: 13th May, 2024
COMPLETED: 13th July, 2024""", )
doc.add_page_break()

# Declaration section formatted
doc.add_heading(text="DECLARATION", level=1)
doc.add_paragraph(text="""I declare that this attachment report is my original work, guided by the three-month attachment period at SAFRA DATA in Nairobi. This report has not been submitted in any form to any other institution of higher learning for the award of an academic qualification. 

I am, therefore, presenting it to Meru University of Science and Technology, School of Pure and Applied Sciences (SPAS), Department of Mathematics, in partial fulfillment for a Bachelor's degree in Mathematics and Computer Science.""")
doc.add_paragraph("\nNAME: MULWA KITALA JEREMIAH.", style='Normal')
doc.add_paragraph("Signature: ................", style='Normal')
doc.add_paragraph("Date: ..........................", style='Normal')
doc.add_page_break()

# Acknowledgement section formatted
doc.add_heading(text="ACKNOWLEDGEMENTS", level=1)
doc.add_paragraph(text="""I would like to express my sincere gratitude to the following individuals and organizations for their invaluable support during my attachment period:""")
doc.add_paragraph(text="- **Mr. Ng'ang'a Ngaruiya**, CEO of SAFRA DATA, for providing me with the opportunity to undertake my attachment at SAFRA DATA and for his insightful guidance and support.", style='List Bullet')
doc.add_paragraph(text="- **My supervisor at SAFRA DATA**, for their continuous mentorship, encouragement, and constructive feedback throughout my attachment.", style='List Bullet')
doc.add_paragraph(text="- **The faculty members of the School of Pure and Applied Sciences (SPAS)** at Meru University of Science and Technology, for their academic support and guidance.", style='List Bullet')
doc.add_paragraph(text="- **Dr. Mutembei**, for coming to assess me as a student at Meru University during my attachment at SAFRA DATA in Nairobi.", style='List Bullet')
doc.add_paragraph(text="- **My family and friends**, for their unwavering support and encouragement throughout my academic journey.", style='List Bullet')
doc.add_paragraph(text="- **My colleagues and peers**, for their collaboration and assistance during the attachment.", style='List Bullet')
doc.add_paragraph(text="- **The respondents in Utawala, Nairobi**, for their willingness to participate in the data collection process and for providing valuable information.", style='List Bullet')
doc.add_page_break()

# Dedication section formatted
doc.add_heading(text="DEDICATION", level=1)
doc.add_paragraph(text="This report is dedicated to:")
doc.add_paragraph(text="- **My parents**, for their unwavering support, encouragement, and sacrifices that have enabled me to pursue my education.", style='List Bullet')
doc.add_paragraph(text="- **My mentors and teachers**, for their guidance, wisdom, and dedication to my academic and personal growth.", style='List Bullet')
doc.add_paragraph(text="- **My friends and peers**, for their camaraderie, support, and the countless moments of shared learning and growth.", style='List Bullet')
doc.add_paragraph(text="- **All those who have inspired me**, for their influence and motivation that have driven me to strive for excellence.", style='List Bullet')
doc.add_page_break()

# Abstract section formatted
doc.add_heading(text="ABSTRACT", level=1)
doc.add_paragraph(text="""This attachment report provides a comprehensive overview of my three-month attachment at SAFRA DATA. The primary objective of the attachment was to gain practical experience in data analysis and to apply theoretical knowledge acquired during my coursework.

During the attachment, I employed a systematic approach to collect data from respondents regarding telecommunications and healthcare. This involved using a structured form accessible via a WhatsApp group link, conducting field visits, and ensuring real-time data entry. Despite some challenges, such as initial reluctance from respondents and occasional technical issues, the data collection process was successful.

The key findings from the data analysis revealed significant insights into the telecommunications and healthcare sectors in Utawala, Nairobi. These insights were valuable to SAFRA DATA and contributed to improving data accuracy and providing actionable recommendations.

In conclusion, the attachment provided me with practical skills in data collection, analysis, and reporting. It also enhanced my ability to handle real-world challenges and improved my professional development. Based on the findings, I recommend continued data collection efforts and further analysis to support decision-making processes at SAFRA DATA.""")
doc.add_page_break()

# Charts section formatted
doc.add_heading(text="CHARTS", level=1)
doc.add_heading(text="1.1. Safra Data Organizational Structure", level=2)
doc.add_picture('C:/Users/mulwa/Desktop/PDFsDOCs/pdfswdoc/SafraOrgStruc.jpg', width=Inches(5.0))
doc.add_page_break()

# Table of Contents placeholder
doc.add_heading(text="TABLE OF CONTENTS", level=1)
doc.add_paragraph(text="Contents will be auto-generated", style='Normal')
doc.add_page_break()

# Methodology section formatted
doc.add_heading(text="METHODOLOGY", level=1)
doc.add_paragraph(text="""During my attachment at SAFRA DATA, I employed a systematic approach to collect data from respondents regarding telecommunications and healthcare. The methodology involved the following steps:

1. **Data Collection Tool**:
   I used a structured form as the primary data collection tool. This form was accessible via a link shared on a WhatsApp group created specifically for the coordination of the attachment program. The form was designed to capture detailed information on various aspects of telecommunications and healthcare.

2. **Respondent Selection**:
   Each day, I aimed to collect data from 5 to 10 respondents. The respondents were selected randomly from the field, ensuring a diverse and representative sample.

3. **Data Collection Process**:
   - **Field Visits**: I conducted field visits to meet with respondents in person. During these visits, I explained the purpose of the data collection and the importance of their participation.
   - **Form Access**: Respondents accessed the form via a link on my phone, which I shared with them during the interaction.
   - **Data Entry**: Respondents filled out the form directly on my phone, providing real-time data entry.

4. **Challenges and Solutions**:
   - **Reluctance to Participate**: Some respondents were initially reluctant to provide their information. To address this, I took the time to explain the purpose of the study in detail, emphasizing the confidentiality and importance of their input. This approach helped in gaining their trust and willingness to participate.
   - **Technical Issues**: Occasionally, there were technical issues with accessing the form. In such cases, I ensured that the form was reloaded or accessed through an alternative link to minimize disruptions.

5. **Data Management**:
   - **Data Storage**: The collected data was securely stored on my phone and periodically backed up to a cloud storage service to prevent data loss.
   - **Data Cleaning**: After the data collection phase, I performed data cleaning to ensure accuracy and completeness. This involved checking for any missing or inconsistent entries and rectifying them as necessary.

6. **Data Reporting**:
   After collecting the data, I sent it back to the organization and provided a detailed report on the findings. This report included insights and recommendations based on the analyzed data.

7. **Tools and Software**:
   - **WhatsApp**: Used for communication and coordination within the attachment program group.
   - **Google Forms**: Utilized for creating and distributing the data collection form.
   - **Python**: Employed for data analysis and visualization, leveraging libraries such as Pandas and Matplotlib.""")
doc.add_page_break()

doc.save(path)
