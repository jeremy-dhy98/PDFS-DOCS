import  docx
from docx.shared import Inches
from  pathlib import  Path

path = Path(Path.home().joinpath("Desktop","CSV","john.docx"))

doc = docx.Document()
doc.add_paragraph("""JOHN KALUKU \nKenya Coast National Polytechnic, \nP.O Box 81220-80100,\nMOMBASA.""").add_run("Email: kalukuj60@gmail.com\nPhone: 0746578734").bold=True
doc.paragraphs[0].runs[0].add_break()

doc.add_heading("Objectives", 0)
doc.add_paragraph("""Seeking  an industrial training opportunity at your company to enhance my skills in Mechanical Engineering  Production while positively contributing to the innovative work of your company.""")
# doc.paragraphs[0].runs[0].add_break()

doc.add_heading("Education", 0)
doc.add_paragraph("(2023-Todate)\nDiploma  in Mechanical Engineering (specializing in Production) at Kenya Coast National Polytechnic.")

doc.add_heading("Experience", 0)
doc.add_paragraph("Academic project Project: Designed and fabricated simple mechanical components including shafts and gears, carried out mini-projects on fabrication and assembly of a small gearbox model using standard workshop procedures.Gained hands-on experience with workshop tools and machinery including lathes, milling machines, drilling machines and welding equipments.")

doc.add_heading("Skills", 0)
doc.add_paragraph("Proficient in operating workshop machines.\nProficient in basic welding and fabrication.\nExposure to CNC basics, measurement and inspection using tools like micrometres and vernier callipers.\nStrong problem solving and Analytical skills. \nExcellent teamwork and collaboration in workshop setting. \nTime management and ability to meet deadlines.")

doc.add_heading("Achievements", 0)
doc.add_paragraph("Successfully completed complex workshop projects involving milling and fabrication of components within set deadlines.\nRecognized by lecturers for excellent teamwork during group technical projects.")

doc.add_heading("ExtraCarricular Activities", 0)
doc.add_paragraph("An active member of the school football team.")

doc.add_heading("References", 0)
doc.add_paragraph("""Mr. Kelvin Ochieng \nH.O.D  Mechanical Engineering \nPhone: 0721814923""")

# img_path =  "IMG_20250208_202537.jpg"
# doc.add_picture(img_path, width=Inches(3), height=Inches(4))
doc.save(path)