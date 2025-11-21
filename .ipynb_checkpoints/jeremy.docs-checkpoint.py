import  docx
from docx.shared import Inches
from  pathlib import  Path

path = Path(Path.home().joinpath("Desktop","CSV","jeremy.docx"))

doc = docx.Document()
doc.add_paragraph("""MULWA KITALA JEREMIAH\nMeru University of Science and \
Technology\nP.O Box 972-60200\nMERU.""").add_run("Email: mulwajeremy24@gmail.com\nPhone: 0713916894").bold=True
doc.paragraphs[0].runs[0].add_break()

doc.add_heading("Objectives", 0)
doc.add_paragraph("""Seeking  an internship opportunity at your company to enhance my skills in Mathematics and Computer Science while positively contributing to the innovative work of your company.""")
# doc.paragraphs[0].runs[0].add_break()

doc.add_heading("Education", 0)
doc.add_paragraph("(2021-Todate)\nBachelor of Science in Mathematics and Computer Science (specializing in Statistics) at Meru University of Science and Technology.")

doc.add_heading("Experience", 0)
doc.add_paragraph("Data Analysis Project: Performed Exploratory Data Analysis(EDA) identifying trends and patterns in data, provided descriptive statistics, and generated visualizations and, reports for decision-making.")

doc.add_heading("Skills", 0)
doc.add_paragraph("Primary data collection and entry.\nProficient with programming languages such as Python and C#.\nStrong Mathematical skills including Calculus, Statistics, and Linear Algebra.\nExperience with MATLAB, SQL, SPSS, STATA, Database management, and Data Analysis tools including Pandas.")

doc.add_heading("Achievements", 0)
doc.add_paragraph("Participated in Kesap Research Center and successfully completed a course in  Statistical Package for Social Sciences(SPSS) and STATA Statistical Package attaining a Certificate.")

doc.add_heading("ExtraCarricular Activities", 0)
doc.add_paragraph("An active member of the Meru Innovation Club.")

doc.add_heading("References", 0)
doc.add_paragraph("""Mrs. Christine Gacheri \nH.O.D Mathematics \nEmail: cmutuura@must.ac.ke \nPhone: 0723674987""")

img_path =  "IMG_20250208_202537.jpg"
doc.add_picture(img_path, width=Inches(3), height=Inches(4))
doc.save(path)