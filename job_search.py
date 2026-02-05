import docx
from docx.shared import Inches, Pt
from pathlib import Path

# --- Configuration ---
# Updated filename for Senior Python Automation Engineer application
path = Path(Path.home().joinpath("Desktop", "CSV", "Jeremiah_Mulwa_Python_Automation_Engineer_CV.docx"))
path.parent.mkdir(parents=True, exist_ok=True)

# Personal details
FULL_NAME = "Jeremiah Mulwa"
ADDRESS_LINES = [
    "Meru, Kenya"
]
EMAIL = "mulwajeremy24@gmail.com"
PHONE = "+254713916894"
GITHUB = "https://github.com/jeremy-dhy98"
LINKEDIN = "https://www.linkedin.com/in/jeremiah-kitala-15aa57288"
PORTFOLIO = "https://github.com/jeremy-dhy98"  # change if you have a portfolio website

# Build document
doc = docx.Document()

# --- Header / Contact ---
header = doc.add_paragraph()
header.alignment = 1  # Center align header
header_run = header.add_run(FULL_NAME + "\n")
header_run.bold = True
header_run.font.size = Pt(18)
header.add_run(f"Email: {EMAIL}    \nPhone: {PHONE}")
header.add_run(f"\nLinkedIn: {LINKEDIN}  \nGitHub: {GITHUB}\n")
header.add_run(f"Portfolio: {PORTFOLIO}\n")

# --- Professional Summary (Tailored for Senior Python Automation Engineer) ---
doc.add_heading("Professional Summary", level=1)
p_summary = doc.add_paragraph()
p_summary.add_run("Senior Python Automation Engineer with hands-on experience designing and owning ").bold = True
p_summary.add_run("scalable automation workflows, API integrations, and data pipelines").bold = True
p_summary.add_run(". Strong focus on system reliability, observability, and production readiness. Experienced in building event-driven systems, implementing robust error handling, and optimizing performance for high-availability services. Skilled at translating business needs into practical, maintainable code and delivering measurable operational improvements.")

# --- Education ---
doc.add_heading("Education", level=1)
p_edu = doc.add_paragraph()
p_edu.add_run("Bachelor of Science in Mathematics and Computer Science").bold = True
p_edu.add_run("\nMeru University of Science and Technology")

# --- Technical Skills (Prioritizing Automation & Systems) ---
doc.add_heading("Technical Skills", level=1)
skills = {
    "Languages & Runtime": "Python (asyncio, multiprocessing, dataclasses, type hints), Bash",
    "Automation & Integration": "APIs (REST), WebSockets, Event-driven design, ETL pipelines",
    "Data & Storage": "Pandas, NumPy, PostgreSQL, SQLite, CSV/JSON handling",
    "Reliability & Ops": "Logging, Structured monitoring, Retries & backoff, Idempotency, Graceful degradation",
    "Containerization & CI": "Docker, CI/CD basics, deployment concepts",
    "Testing & Quality": "Unit & integration testing, mocking external services",
    "Tools & Others": "Git, Docker, FFmpeg, MoviePy (for automation & media tasks)"
}
for k, v in skills.items():
    p = doc.add_paragraph()
    p.add_run(f"{k}: ").bold = True
    p.add_run(v)

# --- Core Experience & Projects (Focused on impact) ---
doc.add_heading("Selected Projects & Experience", level=1)

# Project 1: Automation Pipeline
p1_head = doc.add_paragraph()
p1_head.add_run("End-to-End Automation Pipeline (Event-Driven)").bold = True
doc.add_paragraph("Designed and implemented an event-driven Python pipeline handling real-time data ingestion via WebSockets and REST. Ensured idempotency and used structured logging and monitoring to maintain 99.9% uptime.", style='List Bullet')
doc.add_paragraph("Implemented retry strategies with exponential backoff, circuit-breaker patterns conceptually, and clear failure modes to prevent cascading errors.", style='List Bullet')
doc.add_paragraph("Result: Reduced data processing latency and eliminated duplicate processing, improving operational reliability.", style='List Bullet')

# Project 2: API Integrations & Reliability
p2_head = doc.add_paragraph()
p2_head.add_run("Robust API Integration & System Reliability").bold = True
doc.add_paragraph("Built and maintained multiple API integrations with third-party services, implementing timeouts, retries, and graceful degradation strategies.", style='List Bullet')
doc.add_paragraph("Added structured logging, alerting hooks, and health checks to improve observability and reduce mean time to detect (MTTD).", style='List Bullet')

# Project 3: Data Pipelines & Reporting
p3_head = doc.add_paragraph()
p3_head.add_run("Data Pipelines, Validation & Reporting").bold = True
doc.add_paragraph("Created ETL processes to clean, validate, and transform datasets for downstream analytics using Pandas and SQL.", style='List Bullet')
doc.add_paragraph("Automated scheduled reporting and dashboards to give stakeholders timely, accurate insights.", style='List Bullet')

# Project 4: Automation Tools & Efficiency
p4_head = doc.add_paragraph()
p4_head.add_run("Automation & Scripting for Operational Efficiency").bold = True
doc.add_paragraph("Developed Python and small C# utilities to automate repetitive tasks (file handling, report generation, media processing) improving team productivity.", style='List Bullet')
doc.add_paragraph("Emphasized maintainable code, documentation, and test coverage to ensure long-term reliability.", style='List Bullet')

# --- Achievements & Certifications ---
doc.add_heading("Achievements & Certifications", level=1)
doc.add_paragraph("University Degree: BSc Mathematics & Computer Science — Meru University of Science and Technology", style='List Bullet')
doc.add_paragraph("GitHub portfolio: includes automation pipelines, ETL scripts, and integration examples", style='List Bullet')

# --- References / Availability ---
doc.add_heading("References & Availability", level=1)
doc.add_paragraph("References available upon request. Open to remote, fully-remote, or hybrid opportunities. Available to interview on short notice.", style='List Bullet')

# Save
doc.save(path)
print(f"Professional Senior Python Automation CV saved to: {path}")
