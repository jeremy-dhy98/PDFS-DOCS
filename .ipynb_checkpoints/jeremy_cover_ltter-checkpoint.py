"""
Cover Letter Generator
Generates a professional cover letter (DOCX and plain text) ready to accompany your CV.

Usage examples:
    python Cover_Letter_Generator.py --job_title "Data Analyst" --company "Acme Corp" --recruiter "Hiring Manager"

Or run without args and the script will prompt for missing values interactively.

The script writes the DOCX to ~/Desktop/CSV/cover_letter_{company}_{job_title}.docx by default.

Requires: python-docx (pip install python-docx)
"""

import argparse
from pathlib import Path
import datetime
import textwrap

try:
    import docx
    from docx.shared import Pt
except Exception as e:
    raise SystemExit("Missing dependency: python-docx. Install with 'pip install python-docx'.")

# Default personal details - update if needed
FULL_NAME = "Mulwa Kitala Jeremiah"
EMAIL = "mulwajeremy24@gmail.com"
PHONE = "0713916894"
ADDRESS_LINES = [
    "Meru University of Science and Technology",
    "P.O Box 972-60200",
    "MERU"
]
DEGREE = "B.Sc. Mathematics & Computer Science (Applied Statistics) — 2021-2025"
LINKEDIN = "https://www.linkedin.com/in/jeremiah-kitala-15aa57288"
GITHUB = "https://github.com/jeremy-dhy98"


def render_cover_letter_text(job_title, company, recruiter_name, opening_paragraph, body_paragraphs, closing_paragraph):
    """Return the full cover letter text (string).
    We will include date, contact info, and sign-off.
    """
    today = datetime.date.today().strftime("%B %d, %Y")

    contact = "\n".join(ADDRESS_LINES + [f"Email: {EMAIL}", f"Phone: {PHONE}", f"LinkedIn: {LINKEDIN}", f"GitHub: {GITHUB}"])

    lines = []
    lines.append(contact)
    lines.append("")
    lines.append(today)
    lines.append("")
    # Recipient block
    recipient = recruiter_name if recruiter_name else "Hiring Manager"
    lines.append(f"{recipient}\n{company}\n")

    # Salutation
    lines.append(f"Dear {recipient},")
    lines.append("")

    # Opening paragraph
    lines.append(opening_paragraph)
    lines.append("")

    # Body paragraphs
    for p in body_paragraphs:
        lines.append(p)
        lines.append("")

    # Closing
    lines.append(closing_paragraph)
    lines.append("")
    lines.append("Sincerely,")
    lines.append("")
    lines.append(FULL_NAME)
    lines.append(f"{DEGREE}")

    return "\n".join(lines)


def save_docx(text, out_path: Path):
    doc = docx.Document()
    # Use normal style; set font size for readability
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    for paragraph in text.split('\n'):
        # preserve empty lines
        if paragraph.strip() == "":
            doc.add_paragraph("")
        else:
            doc.add_paragraph(paragraph)

    doc.save(out_path)


def main():
    parser = argparse.ArgumentParser(description="Generate a professional cover letter DOCX and plain-text file.")
    parser.add_argument('--job_title', type=str, help='Job title you are applying for', default='Data Analyst')
    parser.add_argument('--company', type=str, help='Company name', default='[Company Name]')
    parser.add_argument('--recruiter', type=str, help='Name of recruiter or hiring manager', default='Hiring Manager')
    parser.add_argument('--output_dir', type=str, help='Output folder', default=str(Path.home().joinpath('Desktop', 'CSV')))
    parser.add_argument('--personalize', action='store_true', help='Prompt interactively for customizing paragraphs')

    args = parser.parse_args()

    job_title = args.job_title
    company = args.company
    recruiter = args.recruiter

    # Default paragraphs (professional, concise, tailored to your CV)
    opening = (f"I am writing to express my interest in the {job_title} position at {company}. "
               f"I recently graduated with a Bachelor of Science in Mathematics and Computer Science (Applied Statistics) from Meru University of Science and Technology, and I am excited about the opportunity to apply my statistical and programming skills in a data-driven role at {company}.")

    body = [
        ("I am proficient in Python (pandas, numpy), experienced with data manipulation and exploratory data analysis, and familiar with machine learning tools such as scikit-learn and statsmodels. "
         "My strong foundation in statistics and calculus enables me to design and interpret rigorous analyses and build robust models."),

        ("On the engineering side, I have experience building end-to-end data pipelines, fetching historical and real-time data (APIs and WebSockets), and implementing technical indicators and backtesting logic for reproducible analysis. "
         "I am comfortable writing SQL queries and working with relational databases, and I can also contribute software tasks using C# when needed."),

        ("I have completed projects that involved cleaning and transforming large datasets, producing visualizations and reports for decision-making, and automating repetitive tasks to improve efficiency. "
         "I am a proactive learner and collaborate well in teams to deliver production-ready solutions.")
    ]

    closing = (f"I am enthusiastic about the chance to bring analytical rigor and engineering practicality to the {company} team. "
               "I would welcome the opportunity to discuss how my skills and projects align with your needs. Thank you for considering my application.")

    if args.personalize:
        print("Running interactive personalization. Press Enter to accept defaults.\n")
        input_open = input(f"Opening paragraph (default shown):\n{opening}\n\nEnter your custom opening (or press Enter to keep default): ")
        if input_open.strip():
            opening = input_open.strip()

        new_body = []
        for i, b in enumerate(body, start=1):
            print(f"\nBody paragraph {i} (default):\n{b}\n")
            nb = input("Enter custom paragraph (or press Enter to keep default): ")
            new_body.append(nb.strip() if nb.strip() else b)
        body = new_body

        input_close = input(f"\nClosing paragraph (default shown):\n{closing}\n\nEnter custom closing (or press Enter to keep default): ")
        if input_close.strip():
            closing = input_close.strip()

    # Render
    letter_text = render_cover_letter_text(job_title, company, recruiter, opening, body, closing)

    # Prepare output
    out_dir = Path(args.output_dir)
    out_dir.mkdir(parents=True, exist_ok=True)
    safe_company = "_".join(company.split())[:60]
    safe_title = "_".join(job_title.split())[:60]
    filename_base = f"cover_letter_{safe_company}_{safe_title}"

    docx_path = out_dir.joinpath(filename_base + ".docx")
    txt_path = out_dir.joinpath(filename_base + ".txt")

    save_docx(letter_text, docx_path)

    # Save plain text too
    with open(txt_path, 'w', encoding='utf-8') as f:
        f.write(letter_text)

    print(f"Saved cover letter to: {docx_path}")
    print(f"Saved plain-text version to: {txt_path}")
    print("---\nPreview:\n")
    print(textwrap.fill(letter_text, width=80))


if __name__ == '__main__':
    main()
