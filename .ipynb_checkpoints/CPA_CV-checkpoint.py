import docx
from docx.shared import Inches, Pt
from pathlib import Path

# --- Configuration ---
# Updated filename for Clean Power Alliance
path = Path(Path.home().joinpath("Desktop", "CSV", "Jeremiah_Mulwa_CPA_Data_Analyst_CV.docx"))
path.parent.mkdir(parents=True, exist_ok=True)

# Personal details
FULL_NAME = "MULWA KITALA JEREMIAH"
ADDRESS_LINES = [
    "Meru University of Science and Technology",
    "P.O Box 972-60200",
    "MERU"
]
EMAIL = "mulwajeremy24@gmail.com"
PHONE = "0713916894"
GITHUB = "https://github.com/jeremy-dhy98"
LINKEDIN = "https://www.linkedin.com/in/jeremiah-kitala-15aa57288"

# Build document
doc = docx.Document()

# --- Header / Contact ---
header = doc.add_paragraph()
header.alignment = 1 # Center align header
header_run = header.add_run(FULL_NAME + "\n")
header_run.bold = True
header_run.font.size = Pt(18)
header.add_run(f"Email: {EMAIL}\nPhone: {PHONE}\n")
header.add_run(f"LinkedIn: {LINKEDIN}\nGitHub: {GITHUB}\n")

# --- Professional Summary (Tailored for Renewable Energy & CPA) ---
doc.add_heading("Professional Summary", level=1)
p_summary = doc.add_paragraph()
p_summary.add_run("Mission-driven ")
p_summary.add_run("Applied Statistics").bold = True
p_summary.add_run(" graduate with a technical focus on ")
p_summary.add_run("Time-Series Analysis and Predictive Modeling").bold = True
p_summary.add_run(". Passionate about leveraging data science to advance ")
p_summary.add_run("renewable energy solutions and sustainability").bold = True
p_summary.add_run(". Expert in ")
p_summary.add_run("Python (Pandas, Scikit-learn)").bold = True
p_summary.add_run(" for processing large-scale datasets and ")
p_summary.add_run("SQL").bold = True
p_summary.add_run(" for complex data architecture. Committed to supporting Clean Power Alliance’s mission of providing affordable, clean energy through rigorous data integrity and innovative analytical reporting.")

# --- Education ---
doc.add_heading("Education", level=1)
p_edu = doc.add_paragraph()
p_edu.add_run("Bachelor of Science in Mathematics and Computer Science").bold = True
p_edu.add_run("\nSpecialization: Applied Statistics (Expected 2025)")
p_edu.add_run("\nMeru University of Science and Technology")

# --- Relevant Coursework (Refocused for Power/Planning) ---
doc.add_heading("Relevant Coursework", level=2)
courses = [
    "Time Series Analysis & Forecasting (Energy Load Prediction focus)",
    "Statistical Inference and Regression Modeling",
    "Operations Research & Optimization",
    "Database Systems (SQL & Data Warehousing)",
    "Calculus & Mathematical Modeling"
]
for course in courses:
    doc.add_paragraph(course, style='List Bullet')

# --- Technical Skills (Prioritizing Analysis & Engineering) ---
doc.add_heading("Technical Skills", level=1)
skills = {
    "Programming": "Python (NumPy, Pandas, Matplotlib), C#",
    "Data Analysis": "Predictive Modeling, Regression Analysis, Hypothesis Testing, Time-Series Forecasting",
    "Data Engineering": "SQL (PostgreSQL, SQLite), ETL Pipelines, API Integration (REST, WebSockets)",
    "Statistical Tools": "SPSS, STATA, MATLAB, Excel (Advanced Modeling)"
}
for k, v in skills.items():
    p = doc.add_paragraph()
    p.add_run(f"{k}: ").bold = True
    p.add_run(v)

# --- Projects & Experience (Reframed for Utility/Data Impact) ---
doc.add_heading("Projects & Experience", level=1)

# Project 1: Focus on Data Engineering (Relevant to Load Data)
p1_head = doc.add_paragraph()
p1_head.add_run("Automated Data Pipeline & Real-Time Analytics").bold = True
doc.add_paragraph("Developed Python-based ETL pipelines to fetch and process real-time data via WebSockets, ensuring 99.9% data reliability for continuous monitoring.", style='List Bullet')
doc.add_paragraph("Integrated complex logging and error-handling logic to manage large-scale data streams, a skill directly applicable to monitoring utility grid performance.", style='List Bullet')

# Project 2: Focus on Analysis (Relevant to Energy Forecasting)
p2_head = doc.add_paragraph()
p2_head.add_run("Time-Series Analysis & Predictive Reporting").bold = True
doc.add_paragraph("Utilized statistical models to identify trends and anomalies in historical datasets, providing actionable insights for resource planning.", style='List Bullet')
doc.add_paragraph("Cleaned and transformed large-scale datasets using Pandas, preparing data for regression models used in demand forecasting.", style='List Bullet')
doc.add_paragraph("Visualized complex data findings into summary reports for non-technical stakeholders, facilitating evidence-based decision making.", style='List Bullet')

# Project 3: General Skills
p3_head = doc.add_paragraph()
p3_head.add_run("Optimization & Automation Tools").bold = True
doc.add_paragraph("Created custom C# and Python scripts to automate repetitive data extraction tasks, increasing workflow efficiency by approximately 40%.", style='List Bullet')

# --- Achievements & Certifications ---
doc.add_heading("Achievements & Certifications", level=1)
doc.add_paragraph("Professional Certificate: Statistical Analysis using SPSS & STATA", style='List Bullet')
doc.add_paragraph("Innovation Contributor: Meru University Innovation Club (Tech Focus)", style='List Bullet')

# --- References ---
doc.add_heading("References", level=1)
doc.add_paragraph("Mrs. Christine Gacheri-H.O.D Mathematics\nEmail: cmutuura@must.ac.ke\nPhone: 0723674987")

# Save
doc.save(path)
print(f"Perfect CPA CV saved to: {path}")