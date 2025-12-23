"""
CPA Cover Letter Generator
Generates a professional, targeted cover letter for Clean Power Alliance (CPA).

Tailored for Data Analyst / Power Planning roles by emphasizing renewable energy 
objectives, time-series forecasting, and statistical modeling.
"""

import argparse
from pathlib import Path
import datetime
import textwrap

try:
    import docx
    from docx.shared import Pt
except Exception as e:
    raise SystemExit("Missing dependency: python-docx. Install with 'pip install python-docx'.")

# Default personal details
FULL_NAME = "Mulwa Kitala Jeremiah"
EMAIL = "mulwajeremy24@gmail.com"
PHONE = "0713916894"
ADDRESS_LINES = [
    "Meru University of Science and Technology",
    "P.O Box 972-60200",
    "MERU"
]
DEGREE = "B.Sc. Mathematics & Computer Science (Applied Statistics) — 2021-2025"
LINKEDIN = "https://www.linkedin.com/in/jeremiah-kitala-15aa57288"
GITHUB = "https://github.com/jeremy-dhy98"


def render_cover_letter_text(job_title, company, recruiter_name, opening_paragraph, body_paragraphs, closing_paragraph):
    """Return the full cover letter text as a single string."""
    today = datetime.date.today().strftime("%B %d, %Y")
    contact = "\n".join(ADDRESS_LINES + [f"Email: {EMAIL}", "\n" f"Phone: {PHONE}", "\n"f"LinkedIn: {LINKEDIN}", "\n"f"GitHub: {GITHUB}"])

    lines = [
        contact,
        "",
        today,
        "",
        f"{recruiter_name if recruiter_name else 'Hiring Manager'}\n{company}\n",
        f"Dear {recruiter_name if recruiter_name else 'Hiring Manager'},",
        "",
        opening_paragraph,
        ""
    ]
    for p in body_paragraphs:
        lines.append(p)
        lines.append("")

    lines.append(closing_paragraph)
    lines.append("")
    lines.append("Sincerely,")
    lines.append("")
    lines.append(FULL_NAME)
    lines.append(f"{DEGREE}")

    return "\n".join(lines)


def save_docx(text, out_path: Path, job_title, company):
    """Saves the document, applying bolding to key industry terms for CPA."""
    doc = docx.Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Keywords focused on Renewable Energy and Data Science for CPA
    bold_keywords = {
        job_title: job_title,
        company: company,
        "Applied Statistics": "Applied Statistics",
        "renewable energy": "renewable energy",
        "Clean Power Alliance": "Clean Power Alliance",
        "Clean Energy": "Clean Energy",
        "Load Forecasting": "Load Forecasting",
        "Time Series Modeling": "Time Series Modeling",
        "carbon footprint": "carbon footprint",
        "Sustainability": "Sustainability",
        "Python": "Python",
        "SQL": "SQL",
        "ETL": "ETL",
        "Data Engineering": "Data Engineering",
        "predictive insights": "predictive insights",
        "optimization": "optimization"
    }

    sorted_keywords = sorted(bold_keywords.keys(), key=len, reverse=True)

    for paragraph_text in text.split('\n'):
        if not paragraph_text.strip():
            doc.add_paragraph("")
            continue

        p = doc.add_paragraph()
        current_text = paragraph_text

        while current_text:
            found = False
            best_match, best_start, best_end = None, -1, -1

            for keyword in sorted_keywords:
                start_idx = current_text.lower().find(keyword.lower())
                if start_idx != -1 and (best_match is None or start_idx < best_start):
                    best_match, best_start, best_end = keyword, start_idx, start_idx + len(keyword)
                    found = True

            if found:
                if best_start > 0:
                    p.add_run(current_text[:best_start])
                bold_run = p.add_run(current_text[best_start:best_end])
                bold_run.bold = True
                current_text = current_text[best_end:]
            else:
                p.add_run(current_text)
                current_text = ""

    doc.save(out_path)


def main():
    parser = argparse.ArgumentParser(description="Generate a targeted cover letter for Clean Power Alliance.")
    parser.add_argument('--job_title', type=str, default='Data Analyst', help='Targeted job title')
    parser.add_argument('--company', type=str, default='Clean Power Alliance (CPA)', help='Company name')
    parser.add_argument('--recruiter', type=str, default='CPA Recruitment Team', help='Hiring contact')
    parser.add_argument('--output_dir', type=str, default=str(Path.home().joinpath('Desktop', 'CSV')))
    parser.add_argument('--personalize', action='store_true')

    args = parser.parse_args()

    # --- CPA Specific Content ---
    opening = (
        f"I am writing to express my enthusiastic interest in the {args.job_title} position at {args.company}. "
        f"As an Applied Statistics graduate with a deep interest in the renewable energy transition, I am eager to apply my skills in predictive modeling and data engineering to support CPA’s mission of providing clean, affordable electricity to Southern California communities."
    )

    body = [
        ("My background in Applied Statistics has equipped me with a robust understanding of Time Series Modeling and Regression Analysis—tools that are essential for accurate Load Forecasting and energy procurement planning. "
         "I am proficient in using Python (pandas, scikit-learn) and SQL to transform raw utility data into predictive insights that drive operational efficiency and support sustainability goals."),
        
        ("I have practical experience in Data Engineering, specifically building automated ETL pipelines using APIs and WebSockets to handle real-time data streams. "
         "I understand that data integrity is the backbone of utility planning, and I pride myself on developing stable, well-documented systems that ensure accuracy in reporting and carbon footprint tracking."),

        ("Driven by a passion for clean energy, I am committed to the analytical rigor required to navigate the complex regulatory and technical landscape of the energy sector. "
         "I am eager to contribute my technical versatility and proactive problem-solving approach to help Clean Power Alliance optimize its resource portfolio and expand its regional impact.")
    ]

    closing = (
        f"I am highly motivated to join a mission-driven organization like {args.company} where data-driven strategies directly contribute to a greener future. "
        f"Thank you for your time and consideration; I look forward to the possibility of discussing how my statistical foundation can serve your team.")

    if args.personalize:
        # Simple CLI interaction for personalization
        print("Interactive Mode: Personalizing CPA Letter...")
        user_input = input("Enter custom opening (or press enter for default): ")
        if user_input.strip(): opening = user_input.strip()

    letter_text = render_cover_letter_text(args.job_title, args.company, args.recruiter, opening, body, closing)
    
    out_dir = Path(args.output_dir)
    out_dir.mkdir(parents=True, exist_ok=True)
    filename = f"CPA_Cover_Letter_{args.job_title.replace(' ', '_')}.docx"
    save_docx(letter_text, out_dir / filename, args.job_title, args.company)

    print(f"\n[SUCCESS] Generated CPA Cover Letter: {out_dir / filename}")


if __name__ == '__main__':
    main() 