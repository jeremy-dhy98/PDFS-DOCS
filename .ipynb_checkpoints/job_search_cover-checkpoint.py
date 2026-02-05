"""
Senior Python Automation Engineer Cover Letter Generator

Produces a professional, targeted cover letter tailored to Senior Python Automation Engineer
roles (hands-on automation, API integrations, event-driven pipelines, and system reliability).

Usage:
    python generate_cover_letter.py --job_title "Senior Python Automation Engineer" \
        --company "Confidential SaaS / AI-Enabled Platform" --recruiter "Racheal Tutuq" \
        --output_dir "/path/to/output" --personalize
"""

import argparse
from pathlib import Path
import datetime
import textwrap

try:
    import docx
    from docx.shared import Pt
except Exception as e:
    raise SystemExit("Missing dependency: python-docx. Install with 'pip install python-docx'.")


# --- Personal details (edit if needed) ---
FULL_NAME = "Jeremiah Mulwa"
EMAIL = "mulwajeremy24@gmail.com"
PHONE = "+254713916894"
ADDRESS_LINES = [
    "Meru, Kenya"
]
LINKEDIN = "https://www.linkedin.com/in/jeremiah-kitala-15aa57288"
GITHUB = "https://github.com/jeremy-dhy98"
DEGREE = "B.Sc. Mathematics & Computer Science"


def render_cover_letter_text(job_title, company, recruiter_name, opening_paragraph, body_paragraphs, closing_paragraph):
    """Return the full cover letter text as a single string."""
    today = datetime.date.today().strftime("%B %d, %Y")
    contact = "\n".join(ADDRESS_LINES + [f"Email: {EMAIL}", f"Phone: {PHONE}", f"LinkedIn: {LINKEDIN}", f"GitHub: {GITHUB}"])

    lines = [
        contact,
        "",
        today,
        "",
        f"{recruiter_name if recruiter_name else 'Hiring Manager'}\n{company}\n",
        f"Dear {recruiter_name if recruiter_name else 'Hiring Manager'},",
        "",
        opening_paragraph,
        ""
    ]
    for p in body_paragraphs:
        lines.append(p)
        lines.append("")

    lines.append(closing_paragraph)
    lines.append("")
    lines.append("Sincerely,")
    lines.append("")
    lines.append(FULL_NAME)
    lines.append(DEGREE)

    return "\n".join(lines)


def save_docx(text, out_path: Path, job_title, company):
    """Saves the document, applying bolding to key industry terms relevant to automation roles."""
    doc = docx.Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Keywords to emphasize for automation & systems roles
    bold_keywords = {
        job_title: job_title,
        company: company,
        "Senior Python Automation Engineer": "Senior Python Automation Engineer",
        "Python": "Python",
        "automation": "automation",
        "API": "API",
        "APIs": "APIs",
        "WebSockets": "WebSockets",
        "ETL": "ETL",
        "event-driven": "event-driven",
        "reliability": "reliability",
        "idempotency": "idempotency",
        "observability": "observability",
        "monitoring": "monitoring",
        "logging": "logging",
        "Docker": "Docker",
        "CI/CD": "CI/CD",
        "Pandas": "Pandas",
        "PostgreSQL": "PostgreSQL",
        "GitHub": "GitHub",
        "data pipelines": "data pipelines",
        "scalable": "scalable",
        "performance": "performance",
        "system": "system",
        "systems": "systems",
        "integration": "integration",
        "retries": "retries"
    }

    sorted_keywords = sorted(bold_keywords.keys(), key=len, reverse=True)

    for paragraph_text in text.split('\n'):
        if not paragraph_text.strip():
            doc.add_paragraph("")
            continue

        p = doc.add_paragraph()
        current_text = paragraph_text

        while current_text:
            found = False
            best_match, best_start, best_end = None, -1, -1

            for keyword in sorted_keywords:
                start_idx = current_text.lower().find(keyword.lower())
                if start_idx != -1 and (best_match is None or start_idx < best_start):
                    best_match, best_start, best_end = keyword, start_idx, start_idx + len(keyword)
                    found = True

            if found and best_start >= 0:
                if best_start > 0:
                    p.add_run(current_text[:best_start])
                bold_run = p.add_run(current_text[best_start:best_end])
                bold_run.bold = True
                current_text = current_text[best_end:]
            else:
                p.add_run(current_text)
                current_text = ""

    doc.save(out_path)


def main():
    parser = argparse.ArgumentParser(description="Generate a targeted cover letter for a Senior Python Automation Engineer role.")
    parser.add_argument('--job_title', type=str, default='Senior Python Automation Engineer', help='Targeted job title')
    parser.add_argument('--company', type=str, default='SaaS / AI-Enabled Platform', help='Company name')
    parser.add_argument('--recruiter', type=str, default='', help='Hiring contact (optional)')
    parser.add_argument('--output_dir', type=str, default=str(Path.home().joinpath('Desktop', 'CSV')), help='Output directory')
    parser.add_argument('--personalize', action='store_true', help='Enable interactive personalization prompts')

    args = parser.parse_args()

    # --- Role-specific content (tailored for Senior Python Automation Engineer) ---
    opening = (
        f"I am writing to express my interest in the {args.job_title} opportunity at {args.company}. "
        "I bring hands-on experience designing and owning scalable Python-based automation systems, API integrations, and event-driven data pipelines—focused on reliability, observability, and production readiness."
    )

    body = [
        ("In my recent projects I have designed end-to-end automation pipelines that ingest real-time data via WebSockets and REST, "
         "apply validation and transformation (ETL), and push clean outputs to downstream systems. I prioritize idempotent operations, structured logging, and retry strategies with exponential backoff to maintain continuity under failure conditions."),
        
        ("I build integrations with third-party services and internal APIs with clear timeouts, graceful degradation, and monitoring hooks to ensure rapid detection and recovery. "
         "I also emphasize maintainability through testing, documentation, and modular design—so systems can be owned and evolved with confidence."),
        
        ("I enjoy working closely with product and operations teams to translate business needs into pragmatic automation that delivers measurable time savings and operational improvements. "
         "You can review examples of my work and automation projects on my GitHub (linked below). I am excited by the prospect of taking strong ownership of core automation systems at a company that values autonomy and impact.")
    ]

    closing = (
        "Thank you for considering my application. I have attached my CV and a link to my portfolio/GitHub for your review. "
        "I am available to discuss the role at your convenience and can interview on short notice."
    )

    if args.personalize:
        print("Interactive Mode: Personalize the cover letter content.")
        custom_open = input("Custom opening (press Enter to keep default): ").strip()
        if custom_open:
            opening = custom_open
        print("\nYou will be prompted for up to 3 short body bullets. Press Enter to skip or keep defaults.")
        custom_bodies = []
        for i in range(3):
            b = input(f"Body paragraph {i+1} (press Enter to keep default): ").strip()
            if b:
                custom_bodies.append(b)
            else:
                custom_bodies.append(body[i])
        body = custom_bodies
        custom_close = input("Custom closing (press Enter to keep default): ").strip()
        if custom_close:
            closing = custom_close

    letter_text = render_cover_letter_text(args.job_title, args.company, args.recruiter, opening, body, closing)
    
    out_dir = Path(args.output_dir)
    out_dir.mkdir(parents=True, exist_ok=True)
    filename = f"{FULL_NAME.replace(' ', '_')}_Cover_Letter_{args.job_title.replace(' ', '_')}.docx"
    save_docx(letter_text, out_dir / filename, args.job_title, args.company)

    print(f"\n[SUCCESS] Generated Cover Letter: {out_dir / filename}")


if __name__ == '__main__':
    main()
