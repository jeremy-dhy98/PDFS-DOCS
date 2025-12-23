"""
Cover Letter Generator
Generates a professional cover letter (DOCX and plain text) ready to accompany your CV.

Modified to target the Kenya Ports Authority (KPA) Graduate Trainee position.
Uses python-docx formatting (bolding key terms) instead of markdown characters for a cleaner, professional output.
"""

import argparse
from pathlib import Path
import datetime
import textwrap

try:
    import docx
    from docx.shared import Pt
except Exception as e:
    raise SystemExit("Missing dependency: python-docx. Install with 'pip install python-docx'.")

# Default personal details - update if needed
FULL_NAME = "Mulwa Kitala Jeremiah"
EMAIL = "mulwajeremy24@gmail.com"
PHONE = "0713916894"
ADDRESS_LINES = [
    "Meru University of Science and Technology",
    "P.O Box 972-60200",
    "MERU"
]
DEGREE = "B.Sc. Mathematics & Computer Science (Applied Statistics) — 2021-2025"
LINKEDIN = "https://www.linkedin.com/in/jeremiah-kitala-15aa57288"
GITHUB = "https://github.com/jeremy-dhy98"


def render_cover_letter_text(job_title, company, recruiter_name, opening_paragraph, body_paragraphs, closing_paragraph):
    """Return the full cover letter text (string).
    The text returned here is *without* markdown characters, allowing the DOCX function to apply formatting.
    """
    today = datetime.date.today().strftime("%B %d, %Y")

    contact = "\n".join(ADDRESS_LINES + [f"Email: {EMAIL}", f"Phone: {PHONE}", f"LinkedIn: {LINKEDIN}", f"GitHub: {GITHUB}"])

    lines = []
    lines.append(contact)
    lines.append("")
    lines.append(today)
    lines.append("")
    # Recipient block
    recipient = recruiter_name if recruiter_name else "Hiring Manager"
    lines.append(f"{recipient}\n{company}\n")

    # Salutation
    lines.append(f"Dear {recipient},")
    lines.append("")

    # Opening paragraph
    lines.append(opening_paragraph)
    lines.append("")

    # Body paragraphs
    for p in body_paragraphs:
        lines.append(p)
        lines.append("")

    # Closing
    lines.append(closing_paragraph)
    lines.append("")
    lines.append("Sincerely,")
    lines.append("")
    lines.append(FULL_NAME)
    lines.append(f"{DEGREE}")

    return "\n".join(lines)


def save_docx(text, out_path: Path, job_title, company):
    """Saves the document, applying specific formatting (bolding) for keywords."""
    doc = docx.Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Define key terms to bold for maximum impact
    bold_keywords = {
        job_title: job_title,
        company: company,
        "Applied Statistics": "Applied Statistics",
        "Statistical Inference": "Statistical Inference",
        "Regression Analysis": "Regression Analysis",
        "Time Series Modeling": "Time Series Modeling",
        "Python": "Python",
        "pandas": "pandas",
        "numpy": "numpy",
        "scikit-learn": "scikit-learn",
        "SQL": "SQL",
        "data engineering": "data engineering",
        "pipeline reliability": "pipeline reliability",
        "ETL": "ETL",
        "automated reporting": "automated reporting",
        "operational efficiency": "operational efficiency"
    }

    # Process each line and apply formatting where keywords match
    for paragraph_text in text.split('\n'):
        if paragraph_text.strip() == "":
            doc.add_paragraph("")
            continue

        p = doc.add_paragraph()

        # Split paragraph by keyword to apply bolding selectively
        remaining_text = paragraph_text

        # Find all occurrences of keywords and their positions
        runs_to_add = []

        # We need a predictable order for multi-word phrases and single words
        # Sort keywords by length descending to prioritize multi-word phrases
        sorted_keywords = sorted(bold_keywords.keys(), key=len, reverse=True)

        last_idx = 0
        current_text = paragraph_text

        while current_text:
            found = False
            best_match = None
            best_start = -1
            best_end = -1

            # Find the first occurring bold keyword
            for keyword in sorted_keywords:
                # Find keyword case-insensitively
                start_idx = current_text.lower().find(keyword.lower())

                if start_idx != -1 and (best_match is None or start_idx < best_start):
                    best_match = keyword
                    best_start = start_idx
                    best_end = start_idx + len(keyword)
                    found = True

            if found:
                # 1. Add normal text before the keyword
                if best_start > 0:
                    p.add_run(current_text[:best_start])

                # 2. Add the bold keyword (using the original case from the text)
                bold_run = p.add_run(current_text[best_start:best_end])
                bold_run.bold = True

                # 3. Prepare for the next segment
                current_text = current_text[best_end:]
            else:
                # No more keywords found, add the rest of the text and break
                p.add_run(current_text)
                current_text = ""

    doc.save(out_path)


def main():
    parser = argparse.ArgumentParser(description="Generate a professional cover letter DOCX and plain-text file.")
    # --- Custom Defaults for KPA Graduate Trainee Application ---
    parser.add_argument('--job_title', type=str, help='Job title you are applying for', default='Graduate Trainee')
    parser.add_argument('--company', type=str, help='Company name', default='Kenya Ports Authority (KPA)')
    parser.add_argument('--recruiter', type=str, help='Name of recruiter or hiring manager', default='The General Manager Corporate Services')
    parser.add_argument('--output_dir', type=str, help='Output folder', default=str(Path.home().joinpath('Desktop', 'CSV')))
    parser.add_argument('--personalize', action='store_true', help='Prompt interactively for customizing paragraphs')

    # Parse arguments (will use defaults if not provided)
    args = parser.parse_args()

    job_title = args.job_title
    company = args.company
    recruiter = args.recruiter

    # --- Highly Tailored Paragraphs for KPA Data Role (No inline formatting characters) ---
    opening = (
        f"I am writing to express my strong interest in the {job_title} position at the {company}, as advertised on the Careers Portal. "
        f"As a recent graduate with a B.Sc. in Mathematics and Computer Science (Specialization: Applied Statistics), I possess the analytical foundation, data engineering skills, and proactive mindset required to contribute immediately to KPA's strategic goals of enhancing operational efficiency and excellence."
    )

    body = [
        # Focus on Analytical Rigor (Statistics + Data Analysis)
        ("My academic background has provided me with rigorous expertise in Statistical Inference, Regression Analysis, and Time Series Modeling, essential for forecasting, risk assessment, and performance metrics within a complex logistics environment. "
         "I am proficient in applying Python (pandas, numpy, scikit-learn) to conduct in-depth Exploratory Data Analysis (EDA), transforming raw data into actionable insights that can inform operational decisions on throughput and resource allocation."),

        # Focus on Data Engineering and Reliability
        ("I bring hands-on experience in data engineering and pipeline reliability, having built end-to-end systems to fetch and process real-time data using APIs and WebSockets. "
         "This capability ensures data quality and continuous flow, which is crucial for managing dynamic port operations. Furthermore, I am skilled in SQL for robust data querying and managing relational databases."),

        # Focus on Project Implementation and Impact
        ("My project work has centred on developing tools for data extraction, transformation, and automated reporting. For example, I developed scripts that automated repetitive tasks, demonstrating an ability to improve workflow efficiency and reduce manual error—a core requirement for any high-volume administrative setting. I am keen to apply this practical, impact-focused approach to challenges within KPA's various departments.")
    ]

    closing = (
        f"I am highly motivated by the opportunity to apply my quantitative and technical skills to support the vital role the {company} plays in regional trade. "
        f"My detailed CV further outlines my technical proficiencies and projects. Thank you for considering my application; I look forward to the possibility of discussing this exciting Graduate Trainee role."
    )
    # --- End of Custom Paragraphs ---


    if args.personalize:
        print("Running interactive personalization. Press Enter to accept defaults.\n")
        input_open = input(f"Opening paragraph (default shown):\n{opening}\n\nEnter your custom opening (or press Enter to keep default): ")
        if input_open.strip():
            opening = input_open.strip()

        new_body = []
        for i, b in enumerate(body, start=1):
            print(f"\nBody paragraph {i} (default):\n{b}\n")
            nb = input("Enter custom paragraph (or press Enter to keep default): ")
            new_body.append(nb.strip() if nb.strip() else b)
        body = new_body

        input_close = input(f"\nClosing paragraph (default shown):\n{closing}\n\nEnter custom closing (or press Enter to keep default): ")
        if input_close.strip():
            closing = input_close.strip()

    # Render
    letter_text = render_cover_letter_text(job_title, company, recruiter, opening, body, closing)

    # Prepare output
    out_dir = Path(args.output_dir)
    out_dir.mkdir(parents=True, exist_ok=True)
    safe_company = "_".join(company.split())[:60]
    safe_title = "_".join(job_title.split())[:60]
    filename_base = f"cover_letter_{safe_company}_{safe_title}"

    docx_path = out_dir.joinpath(filename_base + ".docx")
    txt_path = out_dir.joinpath(filename_base + ".txt")

    # Save DOCX with formatting applied
    save_docx(letter_text, docx_path, job_title, company)

    # Save plain text too
    with open(txt_path, 'w', encoding='utf-8') as f:
        f.write(letter_text)

    print(f"Saved cover letter to: {docx_path}")
    print(f"Saved plain-text version to: {txt_path}")
    print("---\nPreview:\n")
    print(textwrap.fill(letter_text, width=80))


if __name__ == '__main__':
    main()