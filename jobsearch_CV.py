import docx
from docx.shared import Inches, Pt
from pathlib import Path

# --- Configuration ---
# Updated filename for Senior Python Automation Engineer application
path = Path(Path.home().joinpath("Desktop", "CSV", "Jeremiah_Mulwa_Python_Automation_Engineer_CV.docx"))
path.parent.mkdir(parents=True, exist_ok=True)

# Personal details
FULL_NAME = "Jeremiah Mulwa"
ADDRESS_LINES = [
    "Meru, Kenya"
]
EMAIL = "mulwajeremy24@gmail.com"
PHONE = "+254713916894"
GITHUB = "https://github.com/jeremy-dhy98"
LINKEDIN = "https://www.linkedin.com/in/jeremiah-kitala-15aa57288"
PORTFOLIO = "https://github.com/jeremy-dhy98"  # change if you have a portfolio website

# -------------------------
# Optional metrics (fill with real numbers where available).
# Leave blank to keep non-numeric, impact-focused wording.
# -------------------------
METRICS = {
    "manual_time_saved_pct": "",      # e.g. "40%"
    "uptime_pct": "",                 # e.g. "99.9%"
    "events_per_hour": "",            # e.g. "50k events/hour"
    "latency_reduction_pct": "",      # e.g. "60%"
    "duplicate_reduction_pct": "",    # e.g. "80%"
    "hours_saved_per_week": "",       # e.g. "6 hours/week"
    "cost_savings_estimate": ""       # e.g. "$10k/year"
}

def metric_or(text_if_missing, key, template):
    """Return formatted metric string if present, otherwise fallback text."""
    val = METRICS.get(key, "")
    if val:
        return template.format(val)
    return text_if_missing

# Build document
doc = docx.Document()

# --- Header / Contact ---
header = doc.add_paragraph()
header.alignment = 1  # Center align header
header_run = header.add_run(FULL_NAME + "\n")
header_run.bold = True
header_run.font.size = Pt(18)
header.add_run(f"Email: {EMAIL}    \nPhone: {PHONE}")
header.add_run(f"\nLinkedIn: {LINKEDIN}  \nGitHub: {GITHUB}\n")
header.add_run(f"Portfolio: {PORTFOLIO}\n")

# --- Professional Summary (Tailored for Senior Python Automation Engineer) ---
doc.add_heading("Professional Summary", level=1)
p_summary = doc.add_paragraph()
p_summary.add_run("Senior Python Automation Engineer with hands-on experience designing and owning ").bold = True
p_summary.add_run("scalable automation workflows, API integrations, and data pipelines").bold = True
p_summary.add_run(". Strong focus on system reliability, observability, and production readiness. Experienced in building event-driven systems, implementing robust error handling, and optimizing performance for high-availability services. Skilled at translating business needs into practical, maintainable code and delivering measurable operational improvements. ")

# Add one measurable highlight line if metrics exist
measurable_parts = []
if METRICS["manual_time_saved_pct"]:
    measurable_parts.append(f"Reduced manual processing by {METRICS['manual_time_saved_pct']}")
if METRICS["uptime_pct"]:
    measurable_parts.append(f"Improved system availability to {METRICS['uptime_pct']}")
if METRICS["hours_saved_per_week"]:
    measurable_parts.append(f"Saved ~{METRICS['hours_saved_per_week']} of manual work per week")

if measurable_parts:
    doc.add_paragraph("Key outcomes: " + "; ".join(measurable_parts) + ".")

# --- Education ---
doc.add_heading("Education", level=1)
p_edu = doc.add_paragraph()
p_edu.add_run("Bachelor of Science in Mathematics and Computer Science").bold = True
p_edu.add_run("\nMeru University of Science and Technology")

# --- Technical Skills (Prioritizing Automation & Systems) ---
doc.add_heading("Technical Skills", level=1)
skills = {
    "Languages & Runtime": "Python (asyncio, multiprocessing, dataclasses, type hints), Bash",
    "Automation & Integration": "APIs (REST), WebSockets, Event-driven design, ETL pipelines",
    "Data & Storage": "Pandas, NumPy, PostgreSQL, SQLite, CSV/JSON handling",
    "Reliability & Ops": "Structured logging, monitoring, Retries & exponential backoff, Idempotency, Graceful degradation",
    "Containerization & CI": "Docker, CI/CD basics, deployment concepts",
    "Testing & Quality": "Unit & integration testing, mocking external services, test coverage",
    "Tools & Others": "Git, Docker, FFmpeg, MoviePy (for automation & media tasks)"
}
for k, v in skills.items():
    p = doc.add_paragraph()
    p.add_run(f"{k}: ").bold = True
    p.add_run(v)

# --- Core Experience & Projects (Focused on impact) ---
doc.add_heading("Selected Projects & Experience", level=1)

# Project 1: Automation Pipeline
p1_head = doc.add_paragraph()
p1_head.add_run("End-to-End Automation Pipeline (Event-Driven) — Owner").bold = True
doc.add_paragraph(
    "Owned design and implementation of an event-driven Python pipeline to ingest real-time data via WebSockets and REST, "
    "validate and transform payloads, and persist clean records for downstream consumers.", style='List Bullet')
doc.add_paragraph(
    "Implemented idempotent processing, structured logging, retry strategies with exponential backoff, and clear failure modes to prevent cascading errors.", style='List Bullet')

# Result line with conditional metrics
if METRICS["events_per_hour"] or METRICS["duplicate_reduction_pct"] or METRICS["latency_reduction_pct"]:
    results = []
    if METRICS["events_per_hour"]:
        results.append(f"processed {METRICS['events_per_hour']}")
    if METRICS["duplicate_reduction_pct"]:
        results.append(f"reduced duplicate processing by {METRICS['duplicate_reduction_pct']}")
    if METRICS["latency_reduction_pct"]:
        results.append(f"reduced end-to-end latency by {METRICS['latency_reduction_pct']}")
    doc.add_paragraph("Result: " + "; ".join(results) + ".", style='List Bullet')
else:
    doc.add_paragraph(
        "Result: Improved data consistency and processing reliability, reduced manual intervention, and established repeatable operational patterns.", style='List Bullet')

# Project 2: API Integrations & Reliability
p2_head = doc.add_paragraph()
p2_head.add_run("Robust API Integration & System Reliability").bold = True
doc.add_paragraph(
    "Built and maintained integrations with third-party APIs and internal services applying timeouts, retries, and graceful degradation to limit operational impact from downstream failures.", style='List Bullet')
doc.add_paragraph(
    "Instrumented services with structured logging, health checks, and alerting hooks to improve observability and lower mean time to detect (MTTD).", style='List Bullet')

if METRICS["uptime_pct"] or METRICS["hours_saved_per_week"]:
    r_parts = []
    if METRICS["uptime_pct"]:
        r_parts.append(f"improved availability to {METRICS['uptime_pct']}")
    if METRICS["hours_saved_per_week"]:
        r_parts.append(f"reduced manual ops by ~{METRICS['hours_saved_per_week']} per week")
    doc.add_paragraph("Result: " + "; ".join(r_parts) + ".", style='List Bullet')
else:
    doc.add_paragraph("Result: Increased reliability and reduced operational load through improved error handling and monitoring.", style='List Bullet')

# Project 3: Data Pipelines & Reporting
p3_head = doc.add_paragraph()
p3_head.add_run("Data Pipelines, Validation & Reporting").bold = True
doc.add_paragraph("Created ETL processes to clean, validate, and transform datasets for downstream analytics using Pandas and SQL.", style='List Bullet')
doc.add_paragraph("Automated scheduled reporting and dashboards to give stakeholders timely, accurate insights.", style='List Bullet')

if METRICS["cost_savings_estimate"]:
    doc.add_paragraph(f"Result: Delivered timely reports and workflow automation saving approximately {METRICS['cost_savings_estimate']}.", style='List Bullet')
else:
    doc.add_paragraph("Result: Reduced report preparation time and improved report accuracy for stakeholders, enabling faster decision-making.", style='List Bullet')

# Project 4: Automation Tools & Efficiency
p4_head = doc.add_paragraph()
p4_head.add_run("Automation & Scripting for Operational Efficiency").bold = True
doc.add_paragraph("Developed reusable Python and C# utilities to automate repetitive tasks (file handling, report generation, media processing), increasing team productivity and reducing manual errors.", style='List Bullet')
doc.add_paragraph("Emphasized maintainable code, documentation, and test coverage to ensure long-term reliability and smoother handoffs.", style='List Bullet')

# --- Key Achievements (explicit, measurable outcomes the hiring team asked for) ---
doc.add_heading("Key Achievements", level=1)
# Prefer to include your real metrics here. The text uses METRICS placeholders if provided.
ach1 = metric_or(
    "Designed automation that significantly reduced manual workload across reporting and ingestion tasks.",
    "manual_time_saved_pct",
    "Designed automation that reduced manual processing by {0}."
)
doc.add_paragraph(ach1, style='List Bullet')

ach2 = metric_or(
    "Improved production system reliability and recovery processes through observability and retries.",
    "uptime_pct",
    "Improved production availability to {0} through reliability and observability improvements."
)
doc.add_paragraph(ach2, style='List Bullet')

ach3 = metric_or(
    "Built event-driven pipelines and integrations that improved data consistency and reduced duplicates.",
    "duplicate_reduction_pct",
    "Built event-driven pipelines that reduced duplicate processing by {0}."
)
doc.add_paragraph(ach3, style='List Bullet')

# --- Achievements & Certifications ---
doc.add_heading("Achievements & Certifications", level=1)
doc.add_paragraph("University Degree: BSc Mathematics & Computer Science — Meru University of Science and Technology", style='List Bullet')
doc.add_paragraph("GitHub portfolio: includes automation pipelines, ETL scripts, and integration examples", style='List Bullet')

# --- References / Availability ---
doc.add_heading("References & Availability", level=1)
doc.add_paragraph("References available upon request. Open to remote, fully-remote, or hybrid opportunities. Available to interview on short notice.", style='List Bullet')

# Save
doc.save(path)
print(f"Professional Senior Python Automation CV saved to: {path}")
