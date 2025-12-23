import docx
from docx.shared import Inches
from pathlib import Path

# --- Configuration ---
# Change this path to where you want the CV saved
path = Path(Path.home().joinpath("Desktop", "CSV", "jeremy_KPA_Graduate_Trainee_CV.docx"))
path.parent.mkdir(parents=True, exist_ok=True)

# Personal details (keep these accurate)
FULL_NAME = "MULWA KITALA JEREMIAH"
ADDRESS_LINES = [
    "Meru University of Science and Technology",
    "P.O Box 972-60200",
    "MERU"
]
EMAIL = "mulwajeremy24@gmail.com"
PHONE = "0713916894"
GITHUB = "https://github.com/jeremy-dhy98"
LINKEDIN = "https://www.linkedin.com/in/jeremiah-kitala-15aa57288"

# Photo (optional) — change the filename or set to None to skip
img_path = None # Set to None for safety in a general script

# Build document
doc = docx.Document()

# --- Header / Contact ---
header = doc.add_paragraph()
header_run = header.add_run(FULL_NAME + "\n")
header_run.bold = True
header_run.font.size = docx.shared.Pt(16)
for line in ADDRESS_LINES:
    header.add_run(line + "\n")
header.add_run(f"Email: {EMAIL} | Phone: {PHONE}\n")
header.add_run(f"LinkedIn: {LINKEDIN} | GitHub: {GITHUB}\n") # Simplified contact line


# --- Professional Summary (Targeted for KPA Data Role) ---
doc.add_heading("Professional Summary", level=1)
p_summary = doc.add_paragraph()
p_summary.add_run("Highly analytical and results-oriented ")
p_summary.add_run("Applied Statistics").bold = True
p_summary.add_run(" graduate with a strong foundation in ")
p_summary.add_run("data science, statistical modeling, and database management").bold = True
p_summary.add_run(". Proven ability to build end-to-end data pipelines, perform complex exploratory data analysis (EDA), and translate data insights into clear, actionable reports. "
                   "Proficient in ")
p_summary.add_run("Python").bold = True
p_summary.add_run(" for data manipulation and analysis (pandas, numpy, scikit-learn) and experienced with ")
p_summary.add_run("SQL").bold = True
p_summary.add_run(" for efficient data querying and schema design. "
                   "Eager to apply technical skill and analytical thinking to optimize operational efficiency and inform strategic decision-making within a dynamic maritime and logistics environment like the Kenya Ports Authority (KPA).")


# --- Education ---
doc.add_heading("Education", level=1)
p_edu = doc.add_paragraph("(2021 - 2025)\n")
p_edu.add_run("Bachelor of Science in Mathematics and Computer Science").bold = True
p_edu.add_run(" — Specialization: Applied Statistics\n")
p_edu.add_run("Meru University of Science and Technology")

# --- Relevant Coursework ---
doc.add_heading("Relevant Coursework", level=2)
for course in [
    "Probability & Mathematical Statistics",
    "Statistical Inference and Regression",
    "Time Series Analysis",
    "Database Systems and SQL",
    "Algorithms and Data Structures"
]:
    doc.add_paragraph(course, style='List Bullet')

# --- Technical Skills (Categorized for Scan-ability) ---
doc.add_heading("Technical Skills", level=1)
skills = {
    "Programming": "Python (pandas, numpy), C#",
    "Data Engineering / Databases": "SQL (Querying, Schema Design), ETL Pipelines, API Integration (REST, WebSockets), SQLite, PostgreSQL (familiar)",
    "Data Analysis / ML": "Exploratory Data Analysis (EDA), Statistical Modeling, scikit-learn, statsmodels, matplotlib, seaborn, Time-Series Analysis",
    "Tools / Other": "Git, Jupyter, MATLAB, SPSS, STATA, Microsoft Excel"
}
for k, v in skills.items():
    p = doc.add_paragraph()
    p.add_run(f"{k}: ").bold = True
    p.add_run(v)


# --- Projects & Experience (Reframed for Business Impact) ---
doc.add_heading("Projects & Experience", level=1)

# Project 1: Data engineering
p1_head = doc.add_paragraph()
p1_head.add_run("Automated Data Pipeline & Backtesting").bold = True
p1_head.add_run(" (Data Engineering Focus)")
doc.add_paragraph("Engineered end-to-end data pipelines to reliably fetch historical and real-time data using APIs and WebSockets.", style='List Bullet')
doc.add_paragraph("Ensured data integrity and pipeline stability by implementing robust reconnection logic and comprehensive logging, which is critical for continuous data-flow in operations.", style='List Bullet')
doc.add_paragraph("Utilized Python to integrate technical indicators and run backtesting loops, demonstrating capability in complex data processing and algorithmic application.", style='List Bullet')

# Project 2: Data analysis
p2_head = doc.add_paragraph()
p2_head.add_run("Exploratory Data Analysis (EDA) & Reporting").bold = True
doc.add_paragraph("Conducted in-depth Exploratory Data Analysis (EDA) on large datasets to proactively identify critical trends, anomalies (outliers), and key performance metrics.", style='List Bullet')
doc.add_paragraph("Performed data cleaning and transformation using Python's pandas library to ensure data quality and readiness for modeling.", style='List Bullet')
doc.add_paragraph("Generated clear visual reports and comprehensive summary findings to support managerial decision-making, showcasing the ability to communicate complex findings effectively.", style='List Bullet')

# Project 3: Software / Tools
p3_head = doc.add_paragraph()
p3_head.add_run("Tooling & Automation for Efficiency").bold = True
doc.add_paragraph("Developed utilities and scripts in Python and C# aimed at data extraction, transformation, and automated reporting.", style='List Bullet')
doc.add_paragraph("Automated repetitive data tasks across various processes, resulting in reduced manual error and significant time savings in data preparation.", style='List Bullet')


# --- Achievements & Certifications ---
doc.add_heading("Achievements & Certifications", level=1)
doc.add_paragraph("Certificate — SPSS & STATA course completed at Kesap Research Center", style='List Bullet')
doc.add_paragraph("Active contributor to university innovation initiatives (Meru Innovation Club)", style='List Bullet')

# --- References ---
doc.add_heading("References", level=1)
doc.add_paragraph("Mrs. Christine Gacheri")
doc.add_paragraph("H.O.D Mathematics")
doc.add_paragraph("Email: cmutuura@must.ac.ke")
doc.add_paragraph("Phone: 0723674987")


# Add photo if available (Conditional block, should be safe with img_path = None)
if img_path:
    try:
        doc.add_picture(img_path, width=Inches(1.5))
    except Exception as e:
        doc.add_paragraph("(Photo file not found: ensure 'img_path' points to a valid image if you want a photo included.)")

# Save
doc.save(path)
print(f"Saved updated CV to: {path}")