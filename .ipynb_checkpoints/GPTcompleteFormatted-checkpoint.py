# -*- coding: utf-8 -*-
"""
Created on Sat Aug 24 08:20:42 2024

@author: mulwa
"""

import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path

# Helper functions for formatting
def set_font(run, font_name='Times New Roman', font_size=12, bold=False, italic=False):
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.bold = bold
    run.italic = italic

def set_paragraph_spacing(paragraph, spacing_before=12, spacing_after=12, line_spacing=1.5):
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_before = Pt(spacing_before)
    paragraph_format.space_after = Pt(spacing_after)
    paragraph_format.line_spacing = line_spacing

# Set the file path for the document
path = Path(Path.home().joinpath("Desktop","CSV", "AttachReportCompleteFormatted.docx"))

# Create a new Document
doc = docx.Document()

# Add School Logo
doc.add_picture('C:/Users/mulwa/Desktop/PDFsDOCs/pdfswdoc/schl_logo.jpeg', width=Inches(1.5))
last_paragraph = doc.paragraphs[-1] 
last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Title Page
title = doc.add_heading("MERU UNIVERSITY OF SCIENCE AND TECHNOLOGY", level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

school = doc.add_heading("SCHOOL OF PURE AND APPLIED SCIENCES (SPAS)", level=1)
school.alignment = WD_ALIGN_PARAGRAPH.CENTER

attachment_report = doc.add_heading("ATTACHMENT REPORT", level=1)
attachment_report.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Personal Info
para = doc.add_paragraph()
run = para.add_run("NAME: ")
set_font(run, bold=True)
run = para.add_run("MULWA KITALA JEREMIAH\n")
set_font(run)

run = para.add_run("REG. NO.: ")
set_font(run, bold=True)
run = para.add_run("SC205/103633/20\n")
set_font(run)

run = para.add_run("COURSE: ")
set_font(run, bold=True)
run = para.add_run("Bachelor of Science in Mathematics and Computer Science (Statistics Option)\n")
set_font(run)

run = para.add_run("YEAR: ")
set_font(run, bold=True)
run = para.add_run("III\n")
set_font(run)

run = para.add_run("DEPARTMENT: ")
set_font(run, bold=True)
run = para.add_run("MATHEMATICS\n")
set_font(run)

run = para.add_run("ORGANIZATION: ")
set_font(run, bold=True)
run = para.add_run("SAFRA DATA\n")
set_font(run)

run = para.add_run("C.E.O.: ")
set_font(run, bold=True)
run = para.add_run("NG'ANG'A NGARUIYA\n")
set_font(run)

run = para.add_run("SUPERVISOR: ")
set_font(run, bold=True)
run = para.add_run("MR. BENARD WENJE\n")
set_font(run)

run = para.add_run("DURATION: ")
set_font(run, bold=True)
run = para.add_run("Three Months\n")
set_font(run)

run = para.add_run("PERIOD: ")
set_font(run, bold=True)
run = para.add_run("13th May, 2024 – 13th August, 2024\n")
set_font(run)

set_paragraph_spacing(para, 6, 6)
para.alignment = WD_ALIGN_PARAGRAPH.LEFT
doc.add_page_break()

# Declaration
doc.add_heading("DECLARATION", level=1).bold = True
para = doc.add_paragraph(
    "I declare that this attachment report is my original work, guided by the three-month attachment period at "
    "SAFRA DATA in Nairobi. This report has not been submitted in any form to any other institution of higher "
    "learning for the award of an academic qualification. I am, therefore, presenting it to Meru University of Science "
    "and Technology, School of Pure and Applied Sciences (SPAS), Department of Mathematics, in partial fulfillment "
    "for a Bachelor's degree in Mathematics and Computer Science."
)
set_paragraph_spacing(para, 6, 6)
para.alignment = WD_ALIGN_PARAGRAPH.LEFT

doc.add_paragraph("Name: Mulwa Kitala Jeremiah")
doc.add_paragraph("Signature: _______________")
doc.add_paragraph("Date: _______________")
doc.add_page_break()

# Acknowledgments
doc.add_heading("ACKNOWLEDGMENTS", level=1).bold = True
para = doc.add_paragraph()
run = para.add_run(
    "I would like to express my sincere gratitude to the following individuals and organizations for their invaluable "
    "support during my attachment period:\n- "
)
set_font(run)
run = para.add_run("Mr. Ng'ang'a Ngaruiya")
set_font(run, bold=True)
run = para.add_run(", CEO of SAFRA DATA, for providing me with the opportunity to undertake my attachment "
    "at SAFRA DATA and for his insightful guidance and support.\n- "
)
set_font(run)

run = para.add_run("My supervisor, Mr. Benard Wenje")
set_font(run, bold=True)
run = para.add_run(", for his continuous mentorship, encouragement, and constructive feedback "
    "throughout my attachment.\n- "
)
set_font(run)

run = para.add_run("The faculty members of the School of Pure and Applied Sciences (SPAS)")
set_font(run, bold=True)
run = para.add_run(", for their academic support and guidance.\n- "
)
set_font(run)

run = para.add_run("Dr. Mutembei")
set_font(run, bold=True)
run = para.add_run(", for assessing my work during the attachment period.\n- "
)
set_font(run)

run = para.add_run("My family and friends")
set_font(run, bold=True)
run = para.add_run(", for their unwavering support and encouragement throughout my academic journey.\n- "
)
set_font(run)

run = para.add_run("The respondents in Utawala, Nairobi")
set_font(run, bold=True)
run = para.add_run(", for their willingness to participate in the data collection process."
)
set_font(run)

set_paragraph_spacing(para, 6, 6)
para.alignment = WD_ALIGN_PARAGRAPH.LEFT
doc.add_page_break()

# Dedication
doc.add_heading("DEDICATION", level=1).bold = True
para = doc.add_paragraph()
run = para.add_run("This report is dedicated to:\n")
set_font(run)

run = para.add_run("- My mother Nancy Mulwa and brother Isaac Mulwa")
set_font(run, bold=True)
run = para.add_run(", for their unwavering support, encouragement, and sacrifices "
    "that have enabled me to pursue my education.\n"
)
set_font(run)

run = para.add_run("- My mentors and teachers")
set_font(run, bold=True)
run = para.add_run(", for their guidance, wisdom, and dedication to my academic and personal growth.\n"
)
set_font(run)

run = para.add_run("- My friends and peers")
set_font(run, bold=True)
run = para.add_run(", for their camaraderie and support.\n"
)
set_font(run)

run = para.add_run("- All those who have inspired me")
set_font(run, bold=True)
run = para.add_run(", for motivating me to strive for excellence."
)
set_font(run)

set_paragraph_spacing(para, 6, 6)
para.alignment = WD_ALIGN_PARAGRAPH.LEFT
doc.add_page_break()

# Abstract
doc.add_heading("ABSTRACT", level=1).bold = True
para = doc.add_paragraph(
    "This attachment report provides a comprehensive overview of my three-month attachment at SAFRA DATA. The primary "
    "objective of the attachment was to gain practical experience in data analysis and apply theoretical knowledge. Data "
    "collection focused on telecommunications and healthcare through field visits and online forms.\n\n"
    "Despite challenges like reluctance from some respondents and technical issues, I successfully collected and analyzed "
    "the data. The findings provided insights that contributed to SAFRA DATA's decision-making process."
)
set_paragraph_spacing(para, 6, 6)
para.alignment = WD_ALIGN_PARAGRAPH.LEFT
doc.add_page_break()

# Charts
doc.add_heading("CHARTS", level=1).bold = True
doc.add_heading("1.1. Safra Data Organizational Structure", level=2)
doc.add_picture('C:/Users/mulwa/Desktop/PDFsDOCs/pdfswdoc/SafraOrgStruc.jpg', width=Inches(5.0))
doc.add_page_break()

# Table of Contents placeholder
doc.add_heading("TABLE OF CONTENTS", level=1)
doc.add_paragraph("Contents will be auto-generated", style='Normal')
doc.add_page_break()

# Introduction
doc.add_heading("INTRODUCTION", level=1).bold = True
para = doc.add_paragraph(
    "This report details my three-month attachment at SAFRA DATA, a data analysis company in Nairobi. The primary "
    "objective of this attachment was to bridge the gap between theoretical knowledge and practical experience, focusing "
    "on data collection, analysis, and reporting. The attachment aimed to equip me with practical skills and expose me "
    "to real-world data challenges in the telecommunications and healthcare sectors."
)
set_paragraph_spacing(para, 6, 6)
para.alignment = WD_ALIGN_PARAGRAPH.LEFT
doc.add_page_break()

# Company Profile
doc.add_heading("COMPANY PROFILE", level=1).bold = True

# Background and History
doc.add_heading("Background and History", level=2).bold = True
para = doc.add_paragraph(
    "Founded in 2016, Safra Data is a pioneering data analytics and machine learning company that has consistently delivered "
    "innovative solutions to a wide range of industries. With a strong focus on leveraging cutting-edge technologies, Safra "
    "Data has become a trusted partner for organizations seeking to unlock the power of data-driven insights."
)
set_paragraph_spacing(para, 6, 6)
para.alignment = WD_ALIGN_PARAGRAPH.LEFT
doc.add_page_break()

# Aims
doc.add_heading("Aims", level=2).bold = True
para = doc.add_paragraph(
    "Safra Data focuses on delivering accurate and relevant insights, harnessing science and technology, upholding principles "
    "of security, simplicity, and speed, and enabling faster, smarter, and bolder decisions."
)
set_paragraph_spacing(para, 6, 6)
para.alignment = WD_ALIGN_PARAGRAPH.LEFT
doc.add_page_break()

# Activities Undertaken
doc.add_heading("ACTIVITIES UNDERTAKEN", level=1).bold = True
para = doc.add_paragraph(
    "During my attachment, I was responsible for:\n- Conducting field visits for data collection on telecommunications and "
    "healthcare.\n- Distributing and collecting structured forms via WhatsApp for real-time data entry.\n- Analyzing collected "
    "data using Python and creating visualizations with Pandas and Matplotlib.\n- Preparing reports on findings and recommending "
    "actionable solutions to SAFRA DATA."
)
set_paragraph_spacing(para, 6, 6)
para.alignment = WD_ALIGN_PARAGRAPH.LEFT
doc.add_page_break()

# Skills Gained
doc.add_heading("SKILLS GAINED", level=1).bold = True
para = doc.add_paragraph(
    "Through this attachment, I acquired various skills, including:\n- Data Collection: Designing and administering structured forms.\n"
    "- Data Analysis: Using Python libraries for data manipulation and visualization.\n- Communication: Enhancing interpersonal skills through "
    "respondent interaction.\n- Problem-Solving: Developing solutions to challenges encountered."
)
set_paragraph_spacing(para, 6, 6)
para.alignment = WD_ALIGN_PARAGRAPH.LEFT
doc.add_page_break()

# Methodology
doc.add_heading("METHODOLOGY", level=1).bold = True
para = doc.add_paragraph(
    "The data collection methodology involved using a structured form accessible via WhatsApp, with field visits to engage respondents. "
    "The data was stored securely and analyzed using Python tools."
)
set_paragraph_spacing(para, 6, 6)
para.alignment = WD_ALIGN_PARAGRAPH.LEFT
doc.add_page_break()

# Results and Findings
doc.add_heading("RESULTS AND FINDINGS", level=1).bold = True
para = doc.add_paragraph(
    "Key findings from the analysis indicated a preference for mobile data services over broadband and issues with healthcare accessibility. "
    "Recommendations for SAFRA DATA's client strategies were made based on these insights."
)
set_paragraph_spacing(para, 6, 6)
para.alignment = WD_ALIGN_PARAGRAPH.LEFT
doc.add_page_break()

# Conclusion
doc.add_heading("CONCLUSION", level=1).bold = True
para = doc.add_paragraph(
    "The attachment provided valuable exposure to real-world data analysis and collection, enhancing my technical and professional skills. "
    "The knowledge gained will be instrumental in my future career."
)
set_paragraph_spacing(para, 6, 6)
para.alignment = WD_ALIGN_PARAGRAPH.LEFT
doc.add_page_break()

# Recommendations
doc.add_heading("RECOMMENDATIONS", level=1).bold = True
para = doc.add_paragraph(
    "• SAFRA DATA should continue improving its data collection tools.\n• Meru University should ensure future attachments offer diverse project opportunities.\n"
    "• Continuous data analysis should be encouraged to stay updated with industry trends."
)
set_paragraph_spacing(para, 6, 6)
para.alignment = WD_ALIGN_PARAGRAPH.LEFT
doc.add_page_break()

# References
doc.add_heading("REFERENCES", level=1).bold = True
para = doc.add_paragraph(
    "• SAFRA DATA, Company Overview and Reports.\n• Python Documentation, Pandas, and Matplotlib Libraries."
)
set_paragraph_spacing(para, 6, 6)
para.alignment = WD_ALIGN_PARAGRAPH.LEFT
doc.add_page_break()

# Appendices
doc.add_heading("APPENDICES", level=1).bold = True

# Appendix Entries
appendices = [
    "Appendix 1: Safra Data Organizational Structure",
    "Appendix 2: Sample Data Collection Form",
    "Appendix 3: Field Visit Report"
]

for appendix in appendices:
    doc.add_paragraph(appendix)
    set_font(doc.paragraphs[-1].runs[0], bold=True)

# Save the Document
doc.save(path)
